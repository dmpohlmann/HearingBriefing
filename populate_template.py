"""
Populate DCCEEW_Briefing_OutputTemplate.docx with briefing content.
Strategy: keep cover section, delete everything else, rebuild using
doc.add_paragraph() and doc.add_table() which properly inherit template styles.
"""
from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from copy import deepcopy

BASE = "C:/Users/david/repositories/HearingBriefing"
TEMPLATE = f"{BASE}/DCCEEW_Briefing_OutputTemplate.docx"
OUTPUT = f"{BASE}/ECLC_HIB_003_20251114.docx"

doc = Document(TEMPLATE)

# ── 1. Fill cover metadata table (Table 0) ──

cover = doc.tables[0]
cover_data = {
    0: "Senate Environment and Communications Legislation Committee",
    1: "Environment Protection Reform Bill 2025 and related bills",
    2: "Friday, 14 November 2025",
    3: "Environment Protection Reform Bill 2025 and six related bills",
    4: "Hansard",
}
for row_idx, text in cover_data.items():
    cell = cover.cell(row_idx, 1)
    for p in cell.paragraphs:
        for run in p.runs:
            run.clear()
        if p.runs:
            p.runs[0].text = text
        else:
            p.text = text


# ── 2. Delete all body content from Part A heading onward ──

# Keep paragraphs [0]-[5] (title/cover area) and tables [0]-[1] (cover + AI)
# Delete paragraphs [6] onward and tables [2] onward

# Delete tables in reverse
for ti in range(len(doc.tables) - 1, 1, -1):
    tbl = doc.tables[ti]._tbl
    tbl.getparent().remove(tbl)

# Delete paragraphs in reverse
for pi in range(len(doc.paragraphs) - 1, 5, -1):
    p = doc.paragraphs[pi]._element
    p.getparent().remove(p)


# ── 3. Helpers ──

def add_bullet(p, numId=1, ilvl=0):
    """Add bullet/numbering XML to a paragraph.
    numId=1 → bullet (●), numId=2 → decimal (1., 2., 3.)"""
    pPr = p._element.get_or_add_pPr()
    numPr = pPr.makeelement(qn('w:numPr'), {})
    ilvl_el = numPr.makeelement(qn('w:ilvl'), {qn('w:val'): str(ilvl)})
    numId_el = numPr.makeelement(qn('w:numId'), {qn('w:val'): str(numId)})
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)

def add_heading1(text):
    """Add Heading 1 with split runs: 'Part X:' in Aptos SemiBold + rest."""
    p = doc.add_paragraph(style='Heading 1')
    if ':' in text:
        prefix, rest = text.split(':', 1)
        r1 = p.add_run(prefix + ':')
        r1.font.name = 'Aptos SemiBold'
        p.add_run(' ')
        p.add_run(rest.strip())
    else:
        p.add_run(text)
    return p

def add_list_bullet(text, numId=1, ilvl=0):
    """Add a List Paragraph with bullet/numbering."""
    p = doc.add_paragraph(text, 'List Paragraph')
    add_bullet(p, numId=numId, ilvl=ilvl)
    return p

def make_kv_table(data_rows, style='Table Grid'):
    """Create a 2-column key-value table (like witness/hansard tables)."""
    t = doc.add_table(rows=len(data_rows), cols=2)
    t.style = style
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, (key, val) in enumerate(data_rows):
        # Key cell - bold
        kc = t.cell(ri, 0)
        kp = kc.paragraphs[0]
        kr = kp.add_run(key)
        kr.font.name = 'Aptos SemiBold'
        kr.font.size = Pt(10.5)
        # Value cell
        vc = t.cell(ri, 1)
        vp = vc.paragraphs[0]
        vr = vp.add_run(val)
        vr.font.size = Pt(10.5)
    return t

def make_data_table(headers, rows, header_font_size=Pt(9), cell_font_size=Pt(9)):
    """Create a data table with headers and rows."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for ci, h in enumerate(headers):
        cell = t.cell(0, ci)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.size = header_font_size
        r.font.bold = True
    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = t.cell(ri + 1, ci)
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = cell_font_size
    return t

def add_grouping_note(lines):
    """Add the section grouping note as a bordered callout."""
    # Add as a single-column table to mimic the template's bordered note
    t = doc.add_table(rows=2, cols=1)
    # Header cell
    hc = t.cell(0, 0)
    hp = hc.paragraphs[0]
    hr1 = hp.add_run("Note:")
    hr1.font.bold = True
    hr1.font.size = Pt(9)
    hr2 = hp.add_run("\u2002Section grouping")
    hr2.font.name = 'Aptos SemiBold'
    hr2.font.size = Pt(9)
    # Content cell
    cc = t.cell(1, 0)
    cp = cc.paragraphs[0]
    cr = cp.add_run(lines[0])
    cr.font.size = Pt(9)
    for line in lines[1:]:
        cp = cc.add_paragraph()
        cr = cp.add_run(line)
        cr.font.size = Pt(9)
    return t


# ── 4. Build Part A ──

add_heading1('Part A: Executive briefing')
doc.add_paragraph('Hearing significance', 'Heading 2')

significance = [
    "This was the first \u2013 and, before the March 2026 reporting deadline, the only full-day \u2013 public hearing of the Senate Environment and Communications Legislation Committee\u2019s inquiry into the Environment Protection Reform Bill 2025 and six related bills",
    "Seven distinct witness blocs gave evidence across approximately eight hours, covering the full spectrum of affected interests: the Samuel Review author, conservation NGOs, clean energy, biodiversity and legal experts, a philanthropic foundation, mining industry peaks, property developers, and DCCEEW officials alongside the Threatened Species Scientific Committee chair",
    "The hearing exposed fundamental unresolved disputes \u2013 on the legal effect of the RFA exemption, the breadth of the national interest exemption, and whether the bills can demonstrate species-level protection outcomes \u2013 that DCCEEW must address before the committee reports",
    "Prof. Graeme Samuel AC placed on the public record that he does not support the bills in their current form, despite the government framing them as implementing his 2020 review",
    "The Greens, crossbench, and Coalition each identified distinct lines of attack that, taken together, mean the bills cannot pass the Senate without amendment",
    "DCCEEW officials were unable to name a single animal species that would receive additional protection under the bills \u2013 the most publicly visible exposure of the hearing",
    "The committee will report by 24 March 2026; the government\u2019s stated intention to pass the bills by year-end is now under significant pressure",
]
for b in significance:
    add_list_bullet(b)

doc.add_paragraph('Priority flags', 'Heading 2')
doc.add_paragraph('The following items require SES attention, ranked by urgency.', 'Normal')

doc.add_paragraph('Urgent action', 'Heading 3')
for b in [
    "Senator Henderson placed two high-risk questions on notice: (a) copies of all departmental advices to the minister on unacceptable impacts criteria; and (b) a summary of proposed amendments drafted by the department not yet made public \u2013 responses require Secretary-level clearance and careful legal review before submission (Hansard pp. 71\u201372)",
    "DCCEEW officials were unable to provide assessment pathway usage data when asked \u2013 this was taken on notice and should be provided promptly (Hansard p. 75)",
]:
    add_list_bullet(b)

doc.add_paragraph('Significant risks or exposures', 'Heading 3')
for b in [
    "The \u2018what animal will you stop from going extinct\u2019 exchange (Senator Hanson-Young to DCCEEW officials, Hansard pp. 80\u201381) is now on the public record as a clean attack line \u2013 an evidence-based species outcomes response must be prepared before any further hearing appearance",
    "An unresolved legal dispute exists between Samuel\u2019s position (that NES would govern RFA-covered activities) and the NGO/EJA position (that sections 38 and 42 as drafted preclude this) \u2013 urgent legal advice is required",
    "CME WA confirmed publicly that it had not met the minister since the bills were tabled, despite MCA having engaged the minister and PM\u2019s office \u2013 this engagement asymmetry creates a negative record",
]:
    add_list_bullet(b)

doc.add_paragraph('Legislative or policy signals', 'Heading 3')
for b in [
    "Samuel does not support the national interest exemption in its current form \u2013 described it as creating \u2018abuse of power\u2019 risk and \u2018a conga line of lobbyists\u2019 (Hansard p. 7 approx.)",
    "EJA stated that identified legal issues can be resolved by \u2018simple drafting amendments\u2019, directly undermining the government\u2019s framing that addressing NGO concerns would require reopening the legislative package",
    "The mining sector\u2019s central objection \u2013 subjective unacceptable impacts criteria (\u201837 definitions across 8 pages\u2019) \u2013 requires a clause-level response",
    "Prof. Iain Gordon (TSSC Chair) confirmed the government has not asked TSSC to assess species-level outcomes of the bills, and that TSSC was finalising its own submission",
]:
    add_list_bullet(b)

doc.add_paragraph('Forward preparations', 'Heading 3')
for b in [
    "The TSSC submission \u2013 due within days of the hearing \u2013 will likely become a primary reference point for the committee and requires immediate departmental response on publication",
    "Prepare a worked example of the national interest exemption operating at the end of the assessment pathway, for use at the next hearing",
    "Develop a clearer public explanation of the emissions disclosure provision \u2013 what disclosed emissions data is used for and by whom",
    "AMEC\u2019s NORMs/nuclear actions definition concern is technically plausible and may represent an unintended consequence fixable by targeted amendment \u2013 obtain legal advice",
]:
    add_list_bullet(b)

doc.add_paragraph('Political temperature', 'Heading 2')
for b in [
    "The committee is deeply polarised and the government faces difficult legislative arithmetic in the Senate",
    "Labor senators (Grogan, Ananda-Rajah) are visibly anxious about the bills failing \u2013 Senator Grogan stated she does not want to explain to the environment movement why the legislation failed, while refusing to be held hostage by those demanding perfection",
    "The Greens (Hanson-Young) ran a disciplined, forensic attack on three specific provisions: the RFA exemption, the continuous use exemption, and the absence of a climate trigger \u2013 questioning was well-briefed and designed to build an evidentiary record for Greens amendments",
    "The Coalition (Henderson, McDonald, Duniam) focused on business certainty \u2013 unacceptable impacts criteria, EPA accountability, and inadequate consultation \u2013 and is actively seeking documentation to build a dissenting report",
    "Senator Pocock asked the most technically precise questions, particularly on the offset scheme, the RFA exemption, and emissions disclosure scope \u2013 his alignment with the conservation position on exemptions but independent focus on offset integrity makes him a pivotal negotiating partner",
    "The hearing made clear the bills will require amendment to pass \u2013 the question is whose amendments, on which provisions, and at what political cost",
]:
    add_list_bullet(b)

doc.add_paragraph('Forward look', 'Heading 2')
for b in [
    "The committee will report by 24 March 2026 and the government wants passage before year-end \u2013 these two timelines are now in tension",
    "The critical path runs through three groups simultaneously: the Greens (who need movement on RFA/continuous use exemptions or a climate trigger), the mining sector (who need a cleaner unacceptable impacts definition), and at least a portion of the crossbench",
    "DCCEEW must prepare for the next hearing with: (a) a defensible species-outcomes brief with TSSC input; (b) QoN responses to Henderson\u2019s requests on departmental advices and draft amendments; (c) assessment pathway usage data; (d) a worked example of the national interest exemption in practice; (e) a clear explanation of how emissions disclosure data feeds into decision-making",
    "The TSSC submission must be treated as a priority intelligence item \u2013 it will be the most technically authoritative document the committee receives and will shape committee recommendations on MNES scope and unacceptable impacts",
]:
    add_list_bullet(b)


# ── 5. Build Part B ──

add_heading1('Part B: Section briefings')

add_grouping_note([
    "This hearing is grouped by witness bloc in chronological order of appearance.",
    "Eight distinct blocs appeared across the day, each representing a coherent stakeholder interest.",
    "Grouping by witness bloc best serves analytical clarity because the most significant intelligence derives from the positioning, credibility, and internal consistency of each stakeholder group.",
    "Cross-cutting thematic analysis is provided in Part A and Part C; provision-level cross-referencing is provided in L2 of each section.",
])

# ── Section data ──

sections = [
    {
        "title": "Section 1: Prof. Graeme Samuel AC \u2013 Private capacity (Review author)",
        "hansard": "pp. 1\u201311",
        "witnesses": "Prof. Graeme Samuel AC, private capacity (author, 2020 Independent Review of the EPBC Act)",
        "L1": [
            "Appeared in private capacity as author of the 2020 Independent Review of the EPBC Act",
            "Questioned by Senators Grogan, Hanson-Young, Henderson, Pocock, Walker, and Thorpe",
            "Key topics: national interest exemption \u2013 scope, abuse of power risk; RFA and continuous use exemptions; EPA structure and decision-making role; climate trigger; unacceptable impacts criteria and NES; scope of ministerial discretion",
            "No documents tabled",
            "Acknowledged using AI tools to draft guidance on the national interest definition during the consultative process",
            "Praised Victoria\u2019s early cessation of native forest logging; described a \u2018self-help hierarchy\u2019 where NES rules determine permissibility without ministerial intervention",
        ],
        "L2": [
            "National interest exemption \u2013 does not support the current form; described it as creating an \u2018abuse of power\u2019 risk and a \u2018conga line of lobbyists\u2019; advocated removal or incorporation into the MNES balancing test; diverges from DCCEEW\u2019s public framing (Hansard pp. 5\u20137)",
            "RFA exemption \u2013 stated he \u2018hates\u2019 it but acknowledged political constraints (state agreements to 2030); advocated robust NES to govern RFA conduct; his position that NES would apply is disputed by NGOs on legal grounds (sections 38 and 42)",
            "EPA structure \u2013 prefers audit, oversight, and compliance role only; predicted EPA decision role would diminish to \u2018almost nothing\u2019 once states are accredited; broadly consistent with DCCEEW but more sceptical of EPA decision role",
            "Climate \u2013 opposes a climate trigger; prefers emissions disclosure for projects above 100,000 tonnes; consistent with the bills as drafted",
            "NES specificity \u2013 emphasised granular standards as the mechanism for certainty and reduced litigation",
        ],
        "L3": [
            "Unusual political dynamic: both government and critics claimed Samuel as a supporter",
            "His statement that he does not support the bills in their current form is the most significant political signal from this session \u2013 gives the committee a credible basis for recommending amendments without defeating the bills",
            "Labor senators worked to establish Samuel supports the \u201880 per cent\u2019 approach and that business stakeholders moved the goalposts",
            "Senator Hanson-Young pressed on whether NES would govern RFA conduct \u2013 Samuel maintained they would; NGOs later disputed this",
            "Coalition used Samuel\u2019s evidence selectively \u2013 cited EPA accountability concerns without engaging with his pro-environment positions",
        ],
        "L4": [
            "No formal questions on notice placed in this session",
            "Samuel\u2019s statement that he does not support the bills in current form is now on the public record \u2013 prepare written responses to each of his specific objections",
            "The legal dispute over whether NES applies to RFA-covered activities requires urgent legal advice",
        ],
        "L5": [
            "Samuel is likely to be invited to provide a supplementary submission and may reappear at a subsequent hearing",
            "His position \u2013 supportive of the reform architecture but not current text \u2013 is the most likely vehicle for a bipartisan recommendation on the national interest exemption",
            "DCCEEW should monitor for any supplementary submission and prepare responses",
        ],
    },
    {
        "title": "Section 2: Conservation NGO bloc \u2013 WWF, Greenpeace, ACF",
        "hansard": "pp. 12\u201323",
        "witnesses": "Nicole Forrester, Chief Regenerative Officer, WWF Australia\nEleanor (Elle) Lawless, Senior Nature Campaigner, Greenpeace Australia Pacific\nDavid Ritter, CEO, Greenpeace Australia Pacific\nKelly O\u2019Shanassy, CEO, Australian Conservation Foundation\nBrendan Sydes, National Biodiversity Policy Adviser, ACF",
        "L1": [
            "WWF Australia, Greenpeace Australia Pacific, and Australian Conservation Foundation appeared jointly",
            "Questioned by Senators Hanson-Young, Grogan, Pocock, Henderson, and Ananda-Rajah",
            "Key topics: legal effect of sections 38 and 42 on RFA-covered forests; continuous use exemption and GBR catchment land clearing (500,000 hectares cited); absence of a climate trigger; national interest exemption; offset scheme (\u2018pay to destroy\u2019)",
            "O\u2019Shanassy tabled written materials",
            "Ritter: \u2018Have you ever tried to have a bath where the bath is 95 per cent complete but the plug is not in?\u2019",
            "O\u2019Shanassy: \u2018This is not there. It\u2019s nowhere near there, but it can get there\u2019 \u2013 signalling conditional support",
        ],
        "L2": [
            "RFA exemption (sections 38 and 42) \u2013 NGOs argued NES would not legally govern RFA-covered activities as drafted; material legal dispute unresolved in DCCEEW\u2019s public position",
            "Continuous use exemption \u2013 500,000 hectares cleared in GBR catchment cited; DCCEEW has no public response to this data",
            "Climate trigger \u2013 all three organisations characterised its absence as disqualifying",
            "National interest exemption \u2013 NGOs argued fossil fuel projects should never be eligible",
            "Offset scheme \u2013 characterised restoration contribution fund as \u2018pay to destroy\u2019; cited NSW as a failed precedent",
        ],
        "L3": [
            "Senator Hanson-Young built a formal evidentiary record for Greens amendments on all three exemptions and the climate trigger",
            "Senator Grogan pushed back \u2013 pressing NGOs to acknowledge positive bill elements",
            "NGOs acknowledged positive elements while sustaining core objections \u2013 signals conditional support contingent on amendments",
            "Joint appearance with coordinated arguments signals coalition discipline that will sustain throughout the inquiry",
        ],
        "L4": [
            "No formal questions on notice placed",
            "Exposure: legal dispute on sections 38 and 42 is unresolved \u2013 DCCEEW needs a definitive legal position",
            "Exposure: 500,000 hectare GBR catchment figure was cited without challenge \u2013 verify and prepare response",
        ],
        "L5": [
            "NGO witnesses likely to be called again or provide supplementary submissions",
            "The legal dispute on sections 38 and 42 must be resolved before the committee reports",
            "The climate trigger is the most politically charged issue: without Greens support, passage is uncertain",
        ],
    },
    {
        "title": "Section 3: Clean energy sector \u2013 CEC, Smart Energy Council",
        "hansard": "pp. 24\u201332",
        "witnesses": "William Churchill, Chief Policy and Impact Officer, Clean Energy Council [by video link]\nDavid McElrea, Chief Advocacy Officer, Smart Energy Council [by video link]",
        "L1": [
            "CEC and Smart Energy Council appeared jointly, both by video link",
            "Questioned primarily by Senators Grogan and Ananda-Rajah",
            "Topics: urgency of regulatory reform; offset framework; restoration contribution fund; bioregional planning",
            "No documents tabled; both reserved the right to provide further clarification",
        ],
        "L2": [
            "Standards \u2013 both described NES as \u2018incredibly important\u2019; on-the-record industry endorsement",
            "Offset framework \u2013 strongly supported restoration contribution fund; directly counters the NGO \u2018pay to destroy\u2019 critique",
            "Bioregional planning \u2013 supported as reducing duplication for renewable energy projects",
        ],
        "L3": [
            "Political dynamics: low signal in this section \u2013 evidence was broadly supportive",
            "Government senators used the session to establish on-the-record support from a major economic sector",
        ],
        "L4": [
            "No questions on notice placed; no DCCEEW exposure points identified",
            "Note: both witnesses reserved the right to clarify after member consultation \u2013 monitor for qualified supplementary submissions",
        ],
        "L5": [
            "Clean energy sector likely to be active during Senate debate",
            "Support for restoration contribution fund is the most useful counter-narrative to \u2018pay to destroy\u2019 critique",
        ],
    },
    {
        "title": "Section 4: Biodiversity Council, Environmental Justice Australia, Wilderness Society",
        "hansard": "pp. 33\u201340",
        "witnesses": "James Trezise, CEO, Biodiversity Council\nDr Peter Burnett, Councillor, Biodiversity Council\nProf. Brendan Wintle, Lead Councillor, Biodiversity Council [by video link]\nEllen Maybery, Senior Specialist Lawyer, EJA [by video link]\nNicola Silbert, Senior Lawyer, EJA [by video link]\nSam Szoke-Burke, Biodiversity Policy and Campaign Manager, Wilderness Society",
        "L1": [
            "Biodiversity Council, Environmental Justice Australia, and Wilderness Society appeared jointly",
            "Questioned by Senators Hanson-Young, Pocock, Henderson, Grogan, and Ananda-Rajah",
            "Key topics: offset scheme and \u2018like for like\u2019 test; RFA and continuous use exemptions; \u2018minister satisfied\u2019 drafting test; coal and gas access to streamlined pathway; simple drafting amendments",
            "Wintle provided evidence on the draft offset standard released 24\u201348 hours before hearing",
            "Maybery: identified issues can be resolved by \u2018simple drafting amendments\u2019",
        ],
        "L2": [
            "Offset scheme (Wintle) \u2013 restoration contribution fund enables deviation from \u2018like for like\u2019; drew on NSW experience; called for tightening",
            "\u2018Minister satisfied\u2019 test \u2013 EJA identified it throughout the bills as creating unlimited discretion; requires clause-level response",
            "Coal and gas access \u2013 argued projects can access streamlined pathway with no barrier",
            "Drafting amendments \u2013 Maybery stated fixes are \u2018simple\u2019; directly undermines the government\u2019s framing",
        ],
        "L3": [
            "Senator Pocock\u2019s offset questioning was the most technically sophisticated of the day \u2013 suggests he may pursue an amendment",
            "The EJA \u2018simple drafting amendments\u2019 argument directly undermines the government\u2019s \u2018five years in the making\u2019 framing",
        ],
        "L4": [
            "No formal questions on notice placed",
            "Exposure: draft offset standard released 24\u201348 hours before the hearing \u2013 explain timing if asked",
            "Exposure: Wintle\u2019s NSW offset evidence is specific and factual \u2013 verify or contest",
        ],
        "L5": [
            "The offset scheme is the most technically complex and legally contested provision",
            "Wintle\u2019s evidence will be treated as expert testimony \u2013 prepare a detailed response",
            "The \u2018simple drafting amendments\u2019 argument, if not rebutted, gives the committee a rationale for recommending amendments",
        ],
    },
    {
        "title": "Section 5: Australian Climate and Biodiversity Foundation",
        "hansard": "pp. 41\u201347",
        "witnesses": "Lyndon Schneiders, Executive Director, Australian Climate and Biodiversity Foundation [by video link]",
        "L1": [
            "Lyndon Schneiders appeared by video link",
            "Questioned by Senators Grogan, Henderson, and Hanson-Young",
            "Key topics: overall architecture; net gain; critical habitat; emissions disclosure; EPA independence",
            "Schneiders commended DCCEEW\u2019s work in assembling the legislative package",
            "No documents tabled",
        ],
        "L2": [
            "Overall architecture \u2013 characterised bills as implementing Samuel 2020; identified NES, outcomes in legislation, critical habitat, net gain, and EPA as key strengths",
            "Emissions disclosure \u2013 purpose of disclosed data is unclear; called for clarification of how it relates to climate policy",
        ],
        "L3": [
            "Political dynamics: low signal in this section \u2013 evidence was broadly supportive",
            "Senator Henderson used the session to probe EPA accountability",
        ],
        "L4": [
            "No questions on notice placed; no direct DCCEEW exposure points",
            "Note: emissions disclosure purpose gap is consistent with DCCEEW officials\u2019 afternoon evidence",
        ],
        "L5": [
            "Emissions disclosure purpose gap will likely attract a committee question or recommendation",
            "Develop a clearer public explanation of what decision-makers do with emissions data",
        ],
    },
    {
        "title": "Section 6: Mining industry \u2013 MCA, CME WA, AMEC",
        "hansard": "pp. 48\u201362",
        "witnesses": "Tania Constable, CEO, Minerals Council of Australia\nSteven Brown, Principal Adviser, Environmental Policy, MCA\nChris McCombe, General Manager, Sustainability, MCA\nAnita Logiudice, Director, Policy and Advocacy, CME WA\nWarren Pearce, CEO, Association of Mining and Exploration Companies\nSash Pavic, Director, Commonwealth, AMEC",
        "L1": [
            "MCA, CME WA, and AMEC appeared across the mining industry session",
            "Questioned by Senators Henderson, McDonald, Duniam, Pocock, Grogan, Hanson-Young, and Ananda-Rajah",
            "Key topics: unacceptable impacts criteria (\u201837 definitions across 8 pages\u2019); assessment pathways; EPA accountability; compliance proportionality; climate disclosures; net gain in WA; NORMs/nuclear actions definition",
            "Constable confirmed MCA had met the minister and PM\u2019s office",
            "Logiudice (CME WA) confirmed CME WA had not met the minister since tabling",
            "Pearce (AMEC): \u2018I\u2019d like to argue about something else. I\u2019d really like to see our projects move forward\u2019",
        ],
        "L2": [
            "Unacceptable impacts criteria \u2013 characterised as unclear, subjective, untested; called for single clear definition",
            "EPA accountability \u2013 demanded CEO be removable or EPA limited to compliance/enforcement; aligns with Coalition questioning",
            "Climate disclosures \u2013 sought explicit exclusion from decision-making",
            "Net gain (WA) \u2013 Crown land, remote areas, no baseline data; a genuine implementation gap",
            "NORMs definition \u2013 AMEC argued nuclear actions definition inadvertently captures mineral sands and rare earths; proposed renaming to \u2018radiological exposure actions\u2019",
        ],
        "L3": [
            "Coalition used this session as primary vehicle for building a dissenting report record",
            "Engagement asymmetry between MCA (met minister) and CME WA (not met minister since tabling) is politically significant",
            "AMEC\u2019s appeal to move past argument signals part of the mining sector is separable from MCA/CME WA",
        ],
        "L4": [
            "No formal questions on notice placed",
            "Exposure: CME WA\u2019s public statement requires response \u2013 document departmental engagement",
            "Exposure: verify the 37 definitions claim; prepare explanation if accurate",
            "Exposure: NORMs concern is technically plausible \u2013 obtain legal advice",
        ],
        "L5": [
            "Mining industry witnesses will be a primary source for Coalition dissenting report",
            "Three most likely Coalition amendment vehicles: (a) EPA accountability; (b) unacceptable impacts criteria; (c) NORMs definition",
            "Consider proactively engaging AMEC on NORMs before the next hearing",
        ],
    },
    {
        "title": "Section 7: Property Council of Australia",
        "hansard": "pp. 63\u201367",
        "witnesses": "Matthew Kandelaars, Group Executive, Policy and Advocacy, Property Council of Australia [by video link]\nEleanor Sondergeld, National Policy Manager, Sustainability and Regulatory Affairs, Property Council of Australia [by video link]",
        "L1": [
            "Property Council appeared by video link",
            "Questioned by Senators McDonald, Hanson-Young, Walker, and Grogan",
            "Key topics: August 2025 streamlined approvals for 26,000 homes; urgency of reform (up to four-year wait); bioregional planning; restoration contribution fund",
            "No documents tabled",
        ],
        "L2": [
            "Urgency \u2013 current system \u2018broken\u2019; 26,000 homes streamlined through assessment demonstrated the need",
            "Bioregional planning \u2013 identified as primary mechanism for housing supply improvements",
            "Restoration contribution fund \u2013 supported as enabling pooled offsets and reducing delays",
        ],
        "L3": [
            "Senator Walker built a positive record for reform in the housing context",
            "Property Council evidence is useful for government\u2019s housing narrative but qualified \u2013 real benefit is in elements outside the legislation",
        ],
        "L4": [
            "No questions on notice placed; no direct DCCEEW exposure points",
        ],
        "L5": [
            "Housing narrative depends on elements outside the legislation \u2013 a vulnerability if the committee presses on what the bills themselves deliver",
        ],
    },
    {
        "title": "Section 8: DCCEEW officials and Threatened Species Scientific Committee",
        "hansard": "pp. 68\u201385",
        "witnesses": "Rachel Parry, Deputy Secretary, DCCEEW\nShane Gaddes, Head of Division, Environment Law Reform Taskforce, DCCEEW\nBlaine Wentworth, Acting Branch Head, Policy and Legislation Branch, DCCEEW\nDeclan O\u2019Connor-Cox, Acting Head of Division, Environment Regulation Division, DCCEEW\nGreg Manning, Head of Division, Environment Policy, Regions and Markets, DCCEEW\nAnna-Liisa Lahtinen, Acting Branch Head, Reform Strategy, DCCEEW\nProf. Iain Gordon, Chair, Threatened Species Scientific Committee",
        "L1": [
            "DCCEEW officials appeared for the afternoon session; Prof. Gordon (TSSC Chair) joined slightly late",
            "Questioned by Senators Henderson, Hanson-Young, Grogan, Pocock, McDonald, and Duniam",
            "Key topics: consultation process (minister 100+, department 160+ stakeholders in 6 months); unacceptable impacts criteria; four assessment pathways (accredited, streamlined, EIS, public inquiry); national interest exemption at end of pathway; streamlined pathway (60% uptake, no KPIs, coal/gas access); bilateral accreditation; NOPSEMA; emissions disclosure (scopes 1 and 2); species protection",
            "Senator Hanson-Young: \u2018What animal will you stop from going extinct under this legislation? They can\u2019t answer it.\u2019 (Hansard pp. 80\u201381)",
            "Prof. Gordon: 2,175 species listed; government has not asked TSSC to assess species-level outcomes; TSSC finalising own submission",
            "Senator Henderson placed two QoNs on departmental advices and draft amendments",
            "Bills provided to targeted stakeholders roughly one week before tabling; no NDAs",
        ],
        "L2": [
            "National interest exemption \u2013 sits at END of assessment pathway; full assessment and conditioning required before exemption applies to residual element; \u2018substantial improvement\u2019 on current EPBC exemption",
            "Streamlined pathway \u2013 60% initial uptake; no KPIs; limits public consultation; coal and gas access confirmed",
            "Bilateral accreditation \u2013 states must meet NES, prohibit unacceptable impacts, achieve net gain",
            "Emissions disclosure \u2013 scopes 1 and 2 only; officials could not articulate a decision-making function; scope 3 excluded as \u2018decision of government\u2019",
            "RFA exemption \u2013 no changes in the bills to remove it; removal is \u2018not a simple matter\u2019",
        ],
        "L3": [
            "Henderson\u2019s QoNs signal the Coalition intends to use discovery to examine whether internal advice contradicts public statements",
            "Hanson-Young\u2019s \u2018what animal\u2019 question was a rhetorical trap officials could not escape",
            "Government senators visibly uncomfortable with some answers, particularly on coal/gas access",
            "Senator Pocock pursued targeted questions on emissions scope, RFA exemption, and stakeholder early access",
        ],
        "L4": [
            "QoN 1 \u2013 Henderson: copies of all departmental advices on unacceptable impacts criteria; outstanding",
            "QoN 2 \u2013 Henderson: summary of proposed amendments not yet public; outstanding",
            "QoN 3 \u2013 Multiple senators: assessment pathway usage data; outstanding",
            "QoN 4 \u2013 Pocock: stakeholder early access list, dates, and advice; outstanding",
            "QoN 5 \u2013 Pocock: policy reasons for not removing RFA exemption; outstanding",
            "Exposure (high): the \u2018what animal\u2019 exchange is the most visible public exposure point \u2013 requires species-outcomes response",
            "Exposure (medium): no KPIs for streamlined pathway uptake",
            "Exposure (medium): emissions disclosure purpose gap unresolved",
        ],
        "L5": [
            "QoN responses are the most immediate obligation; Henderson QoNs 1 and 2 require Secretary-level clearance",
            "TSSC submission will be the most consequential external document \u2013 obtain immediately and prepare response",
            "Before the next hearing prepare: worked example of national interest exemption; emissions disclosure explanation; all QoN responses; species-outcomes brief; assessment pathway data",
        ],
    },
]

# Build sections with continuous numbering (numId=2 = decimal, continues across all paragraphs)
for sec in sections:
    doc.add_paragraph(sec["title"], 'Heading 2')
    make_kv_table([
        ("Hansard pages", sec["hansard"]),
        ("Witnesses", sec["witnesses"]),
    ])

    for layer_key, layer_title in [
        ("L1", "L1 \u2013 The record"),
        ("L2", "L2 \u2013 Legislative and policy signal"),
        ("L3", "L3 \u2013 Political intelligence"),
        ("L4", "L4 \u2013 Risk and exposure"),
        ("L5", "L5 \u2013 Forward look"),
    ]:
        doc.add_paragraph(layer_title, 'Heading 3')
        for bullet in sec[layer_key]:
            add_list_bullet(bullet, numId=2)


# ── 6. Build Part C ──

add_heading1('Part C: Quick reference')

doc.add_paragraph('Table 1: Commitments and questions on notice', 'Heading 3')
make_data_table(
    ["No.", "Question or commitment", "Witness", "Senator", "Due date/status"],
    [
        ["1", "Copies of all departmental advices to the minister on unacceptable impacts criteria", "DCCEEW officials (Parry/Gaddes)", "Henderson", "Outstanding \u2013 standard Senate timeline"],
        ["2", "Summary of proposed amendments drafted by the department not yet made public", "DCCEEW officials (Parry)", "Henderson", "Outstanding \u2013 standard Senate timeline"],
        ["3", "Breakdown of most-used current EPBC assessment pathways and proportions (since 1 July 2022)", "DCCEEW officials (O\u2019Connor-Cox/Gaddes)", "Multiple", "Outstanding \u2013 could not be provided on the day"],
        ["4", "List of stakeholders who received early access to the bills, exact dates, and related advice", "DCCEEW officials (Gaddes)", "Pocock", "Outstanding \u2013 standard Senate timeline"],
        ["5", "Policy reasons for not removing the RFA exemption in this legislative package", "DCCEEW officials (Manning/Gaddes)", "Pocock", "Outstanding \u2013 standard Senate timeline"],
    ],
)

doc.add_paragraph('Table 2: Key quotes', 'Heading 3')
make_data_table(
    ["Speaker", "Quote (verbatim \u2013 flag if paraphrased)", "Significance", "Hansard pages"],
    [
        ["Prof. Graeme Samuel AC", "\u2018Abuse of power by a minister. And, also, the risk for the minister is there will be a conga line of lobbyists outside their door.\u2019", "Review author does not support the exemption in current form; credible basis for amendment", "p. 7 approx."],
        ["Kelly O\u2019Shanassy, ACF", "\u2018This is not there. It\u2019s nowhere near there, but it can get there.\u2019", "Signals conditional NGO support \u2013 bills salvageable with amendments", "p. 22 approx."],
        ["David Ritter, CEO, Greenpeace", "\u2018Have you ever tried to have a bath where the bath is 95 per cent complete but the plug is not in?\u2019", "Memorable framing of exemptions problem; likely quoted in committee report", "p. 26 approx."],
        ["Ellen Maybery, EJA", "\u2018Those fixes are clear, can be done, and we call on the Senate to do so.\u2019", "Undermines government framing that fixes require reopening the package", "p. 33 approx."],
        ["Warren Pearce, CEO, AMEC", "\u2018I\u2019d like to argue about something else. I\u2019d really like to see our projects move forward.\u2019", "Part of mining sector prioritises speed over amendment demands", "p. 57 approx."],
        ["Senator Grogan (ALP)", "\u2018I don\u2019t know how you or various people are going to feel if this doesn\u2019t happen, but I know how I\u2019m going to feel.\u2019", "Government senator signalling anxiety about failure", "p. 31 approx."],
        ["Senator Hanson-Young", "\u2018What animal will you stop from going extinct under this legislation? They can\u2019t answer it.\u2019", "Communications vulnerability; will recur at subsequent hearings", "p. 81 approx."],
    ],
)

doc.add_paragraph('Table 3: Forward action items for DCCEEW', 'Heading 3')
make_data_table(
    ["Priority", "Action item", "Suggested owner", "Timeframe"],
    [
        ["HIGH", "Prepare QoN response on departmental advices re unacceptable impacts criteria (Henderson QoN 1) \u2013 requires Secretary-level clearance", "Deputy Secretary / General Counsel", "Immediate \u2013 within 2 weeks"],
        ["HIGH", "Prepare QoN response on undisclosed draft amendments (Henderson QoN 2)", "Deputy Secretary / Minister\u2019s office", "Immediate \u2013 within 2 weeks"],
        ["HIGH", "Prepare assessment pathway usage data table since 1 July 2022 (QoN 3)", "Assessment policy branch", "Within 1 week"],
        ["HIGH", "Prepare QoN responses on stakeholder early access and RFA exemption policy reasons (QoNs 4 and 5)", "Environment Law Reform Taskforce", "Within 2 weeks"],
        ["HIGH", "Obtain TSSC submission on publication; prepare departmental response", "Biodiversity policy / Legal", "Within 1 week of publication"],
        ["HIGH", "Develop species-outcomes brief with specific species and protective mechanism analysis; engage TSSC", "Assessment policy / TSSC liaison", "Before next hearing"],
        ["MEDIUM", "Obtain legal advice on sections 38 and 42 \u2013 do NES apply to RFA-covered activities as drafted?", "Legal team", "Within 2 weeks"],
        ["MEDIUM", "Obtain legal advice on NORMs/nuclear actions definition", "Legal team", "Within 2 weeks"],
        ["MEDIUM", "Develop public explanation of emissions disclosure provision", "Assessment policy / Climate policy", "Before next hearing"],
        ["MEDIUM", "Develop worked example of the national interest exemption at end of assessment pathway", "Assessment policy", "Before next hearing"],
        ["MEDIUM", "Arrange ministerial engagement with CME WA", "Ministerial office", "Within 2 weeks"],
        ["LOW", "Verify \u201837 definitions across 8 pages\u2019 claim; prepare clause-level rebuttal if inaccurate", "Assessment policy / Legal", "Before next hearing"],
        ["LOW", "Verify 500,000 hectare GBR catchment land clearing data; prepare response", "Environment policy branch", "Before next hearing"],
    ],
)


# ── 7. Save ──

doc.save(OUTPUT)
print(f"Saved to {OUTPUT}")
